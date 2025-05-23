"""File processors for different file types."""

import json
import csv
from pathlib import Path
from typing import Dict, List, Optional, Any
from abc import ABC, abstractmethod

import PyPDF2
from openpyxl import load_workbook
from docx import Document


class FileProcessor(ABC):
    """Abstract base class for file processors."""
    
    @abstractmethod
    def can_process(self, file_path: Path) -> bool:
        """Check if this processor can handle the file."""
        pass
    
    @abstractmethod
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content and metadata from the file."""
        pass


class PDFProcessor(FileProcessor):
    """Processor for PDF files."""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "content": text,
                    "page_count": len(reader.pages),
                    "metadata": reader.metadata or {}
                }
        except Exception as e:
            return {"content": "", "error": str(e)}


class ExcelProcessor(FileProcessor):
    """Processor for Excel files."""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".xlsx", ".xls"]
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        try:
            workbook = load_workbook(file_path, data_only=True)
            content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        sheet_data.append([str(cell) if cell is not None else "" for cell in row])
                
                if sheet_data:
                    content.append(f"Sheet: {sheet_name}\n" + "\n".join(["\t".join(row) for row in sheet_data]))
            
            return {
                "content": "\n\n".join(content),
                "sheet_count": len(workbook.sheetnames),
                "sheets": workbook.sheetnames
            }
        except Exception as e:
            return {"content": "", "error": str(e)}


class WordProcessor(FileProcessor):
    """Processor for Word documents."""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".docx", ".doc"]
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        try:
            if file_path.suffix.lower() == ".docx":
                doc = Document(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                return {
                    "content": text,
                    "paragraph_count": len(doc.paragraphs)
                }
            else:
                return {"content": "", "error": "DOC format not supported, only DOCX"}
        except Exception as e:
            return {"content": "", "error": str(e)}


class CSVProcessor(FileProcessor):
    """Processor for CSV files."""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".csv"
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                rows = list(reader)
                
                if rows:
                    headers = rows[0]
                    content = f"Headers: {', '.join(headers)}\n\n"
                    content += "\n".join(["\t".join(row) for row in rows[:100]])  # Limit to first 100 rows
                    
                    return {
                        "content": content,
                        "row_count": len(rows),
                        "column_count": len(headers) if headers else 0,
                        "headers": headers
                    }
                else:
                    return {"content": "", "row_count": 0}
        except Exception as e:
            return {"content": "", "error": str(e)}


class TextProcessor(FileProcessor):
    """Processor for text files."""
    
    def can_process(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".txt", ".md", ".py", ".js", ".ts", ".html", ".css", ".yaml", ".yml", ".json", ".xml", ".toml", ".ini"]
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                return {
                    "content": content,
                    "line_count": len(content.splitlines()),
                    "character_count": len(content)
                }
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read()
                    return {
                        "content": content,
                        "line_count": len(content.splitlines()),
                        "character_count": len(content),
                        "encoding": "latin-1"
                    }
            except Exception as e:
                return {"content": "", "error": str(e)}
        except Exception as e:
            return {"content": "", "error": str(e)}


class FileProcessorManager:
    """Manager for all file processors."""
    
    def __init__(self):
        self.processors = [
            PDFProcessor(),
            ExcelProcessor(),
            WordProcessor(),
            CSVProcessor(),
            TextProcessor()
        ]
    
    def get_processor(self, file_path: Path) -> Optional[FileProcessor]:
        """Get the appropriate processor for a file."""
        for processor in self.processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    def process_file(self, file_path: Path) -> Dict[str, Any]:
        """Process a file and extract its content."""
        processor = self.get_processor(file_path)
        if processor:
            result = processor.extract_content(file_path)
            result["file_path"] = str(file_path)
            result["file_name"] = file_path.name
            result["file_size"] = file_path.stat().st_size
            result["modified_time"] = file_path.stat().st_mtime
            return result
        else:
            return {
                "content": "",
                "error": "No processor available for this file type",
                "file_path": str(file_path),
                "file_name": file_path.name
            }